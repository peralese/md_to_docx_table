# md_to_docx_table.py
# pip install python-docx markdown beautifulsoup4

import argparse, os, sys
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from bs4 import BeautifulSoup

# ---------- helpers: tables ----------

def _write_table(doc: Document, header, rows):
    n_cols = max(len(header), max((len(r) for r in rows), default=0))
    if n_cols == 0:
        return
    table = doc.add_table(rows=(1 if header else 0) + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    def put(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run("" if text is None else str(text))
        run.font.bold = bold
        p.alignment = align
        p.space_before = Pt(0); p.space_after = Pt(0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    if header:
        for j in range(n_cols):
            put(table.rows[0].cells[j], header[j] if j < len(header) else "",
                bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    start_i = 1 if header else 0
    for i, r in enumerate(rows, start=start_i):
        for j in range(n_cols):
            put(table.rows[i].cells[j], r[j] if j < len(r) else "")

def _normalize_pipe_tables(md_text: str) -> str:
    """
    (1) Insert a blank line before any pipe-style table so 'tables' parses it.
    (2) Collapse extra blank lines *inside* a table block (Markdown tables don't allow them).
    """
    lines = md_text.splitlines()
    out, i = [], 0
    def is_pipe_row(s: str) -> bool:
        s = s.strip()
        return s.startswith("|") and s.count("|") >= 2

    while i < len(lines):
        line = lines[i]
        if is_pipe_row(line):
            if out and out[-1].strip() != "":
                out.append("")  # ensure a blank line before a table block
            # consume contiguous table rows, skipping blank lines inside the block
            j = i
            while j < len(lines):
                s = lines[j]
                if s.strip() == "":
                    if j + 1 < len(lines) and is_pipe_row(lines[j + 1]):
                        j += 1
                        continue
                    else:
                        break
                if not is_pipe_row(s):
                    break
                out.append(s)
                j += 1
            i = j
            continue
        out.append(line)
        i += 1
    return "\n".join(out)

def _parse_pipe_table(md_lines, start):
    """Parse a pipe table allowing blank lines between rows. Returns (end_idx, header, rows) or (start, None, None)."""
    rows, i = [], start
    def is_pipe_row(s: str) -> bool:
        s = s.strip()
        return s.startswith("|") and s.count("|") >= 2

    while i < len(md_lines):
        s = md_lines[i].strip()
        if s == "":
            if i + 1 < len(md_lines) and is_pipe_row(md_lines[i + 1]):
                i += 1
                continue
            else:
                break
        if not is_pipe_row(s):
            break
        rows.append([c.strip() for c in s.strip("|").split("|")])
        i += 1

    if len(rows) < 2:
        return start, None, None

    sep = rows[1]
    def is_sep(cell: str) -> bool:
        return set(cell.replace(":", "").replace("-", "")) == set()
    ok_sep = all(is_sep(c) and set(c) <= set("-:") and len(c.replace(":", "")) >= 3 for c in sep)
    if not ok_sep:
        return start, None, None

    header = rows[0]
    body = rows[2:] if len(rows) > 2 else []
    return i, header, body

# ---------- conversion ----------

def md_to_docx(md_text: str, output_path: str):
    # Make tables reliable before HTML conversion
    md_text = _normalize_pipe_tables(md_text)

    html = markdown.markdown(md_text, extensions=["tables", "fenced_code", "attr_list"])
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()
    try:
        font = doc.styles["Normal"].font
        font.name = "Calibri"; font.size = Pt(11)
    except Exception:
        pass

    emitted_table = False
    for elem in soup.contents:
        if elem.name is None:
            t = str(elem).strip()
            if t: doc.add_paragraph(t)
            continue

        if elem.name in ["h1","h2","h3","h4","h5","h6"]:
            p = doc.add_paragraph(elem.get_text())
            try: p.style = f"Heading {min(int(elem.name[1]), 3)}"
            except Exception: pass

        elif elem.name == "p":
            doc.add_paragraph(elem.get_text())

        elif elem.name == "pre" or (elem.name == "code" and elem.parent.name != "p"):
            p = doc.add_paragraph()
            run = p.add_run(elem.get_text())
            run.font.name = "Consolas"; run.font.size = Pt(10.5)

        elif elem.name == "ul":
            for li in elem.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(), style="List Bullet")

        elif elem.name == "ol":
            for li in elem.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(), style="List Number")

        elif elem.name == "table":
            header = []
            thead = elem.find("thead")
            if thead:
                header = [th.get_text(strip=True) for th in thead.find_all("th")]
            body_rows = []
            tbody = elem.find("tbody") or elem
            for i, tr in enumerate(tbody.find_all("tr", recursive=False)):
                cells = tr.find_all(["td","th"], recursive=False)
                texts = [c.get_text(strip=True) for c in cells]
                if i == 0 and not header and any(c.name == "th" for c in cells):
                    header = texts; continue
                body_rows.append(texts)
            _write_table(doc, header, body_rows)
            emitted_table = True

        else:
            doc.add_paragraph(elem.get_text())

    # Very defensive fallback if no HTML tables were emitted at all
    if not emitted_table:
        lines, i = md_text.splitlines(), 0
        while i < len(lines):
            line = lines[i]
            if line.strip().startswith("|") and line.count("|") >= 2:
                j, header, body = _parse_pipe_table(lines, i)
                if header is not None:
                    _write_table(doc, header, body)
                    i = j; continue
            doc.add_paragraph(line); i += 1

    doc.save(output_path)

# ---------- CLI ----------

def main():
    ap = argparse.ArgumentParser(
        prog="md_to_docx_table",
        description="Convert Markdown to DOCX with robust table handling."
    )
    ap.add_argument("input", help="Path to input .md file")
    ap.add_argument("-o", "--output", help="Path to output .docx file (default: alongside input)")
    ap.add_argument("--out-dir", help="Directory for the output .docx (overrides default location)")
    ap.add_argument("-f", "--force", action="store_true", help="Overwrite output if it exists")
    args = ap.parse_args()

    inp = os.path.abspath(args.input)
    if not os.path.isfile(inp):
        print(f"Not found: {inp}"); sys.exit(1)

    base, _ = os.path.splitext(inp)
    outp = args.output or (base + ".docx")
    if args.out_dir:
        os.makedirs(args.out_dir, exist_ok=True)
        outp = os.path.join(os.path.abspath(args.out_dir), os.path.basename(base) + ".docx")

    if os.path.exists(outp) and not args.force:
        print(f"Refusing to overwrite existing file: {outp}\nUse -f/--force to overwrite.")
        sys.exit(1)

    with open(inp, "r", encoding="utf-8") as f:
        md_text = f.read()

    md_to_docx(md_text, outp)
    print(f"Converted {inp} -> {outp}")

if __name__ == "__main__":
    main()

